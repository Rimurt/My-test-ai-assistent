"""
modules/file_reader.py

Универсальный модуль для чтения содержимого файлов различных форматов.
Возвращает текст, пригодный для передачи в языковую модель.
"""

import os
from pathlib import Path
from typing import Tuple, Optional

# Максимальное количество символов, передаваемых в модель (защита от огромных файлов)
MAX_CONTENT_LENGTH = 30_000


def read_file_content(file_path: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Читает файл и возвращает его текстовое содержимое.

    Returns:
        (content, error) — одно из двух будет None.
        content — строка с содержимым файла (или None при ошибке).
        error   — строка с описанием ошибки (или None при успехе).
    """
    path = Path(file_path)
    if not path.exists():
        return None, f"Файл не найден: {file_path}"

    ext = path.suffix.lower().lstrip(".")

    try:
        if ext in ("txt", "md", "markdown", "rst", "log",
                   "json", "xml", "yaml", "yml", "toml",
                   "ini", "cfg", "bat", "sh",
                   "py", "js", "ts", "html", "css",
                   "java", "cpp", "c", "h", "cs",
                   "go", "rs", "php", "rb"):
            return _read_text(path)

        elif ext == "csv":
            return _read_csv(path)

        elif ext == "tsv":
            return _read_csv(path, sep="\t")

        elif ext in ("xlsx", "xls", "ods"):
            return _read_spreadsheet(path)

        elif ext in ("docx", "doc", "odt"):
            return _read_word(path)

        elif ext == "pdf":
            return _read_pdf(path)

        else:
            # Попытка прочитать как обычный текст
            return _read_text(path)

    except Exception as e:
        return None, str(e)


# ─────────────────────────────────────────────
# Внутренние читатели
# ─────────────────────────────────────────────

def _truncate(text: str) -> str:
    if len(text) > MAX_CONTENT_LENGTH:
        cut = text[:MAX_CONTENT_LENGTH]
        return cut + f"\n\n[... файл обрезан: показаны первые {MAX_CONTENT_LENGTH} символов из {len(text)} ...]"
    return text


def _read_text(path: Path) -> Tuple[Optional[str], Optional[str]]:
    """Читает plain-text файл с автоопределением кодировки."""
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            text = path.read_text(encoding=encoding)
            return _truncate(text), None
        except UnicodeDecodeError:
            continue
    return None, "Не удалось определить кодировку файла."


def _read_csv(path: Path, sep: str = ",") -> Tuple[Optional[str], Optional[str]]:
    """Читает CSV/TSV и форматирует как markdown-таблицу."""
    import csv

    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            with open(path, newline="", encoding=encoding) as f:
                reader = csv.reader(f, delimiter=sep)
                rows = list(reader)
            break
        except UnicodeDecodeError:
            continue
    else:
        return None, "Не удалось определить кодировку CSV-файла."

    if not rows:
        return "(файл пуст)", None

    # Формируем markdown-таблицу
    lines = []
    header = rows[0]
    lines.append("| " + " | ".join(str(c) for c in header) + " |")
    lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    for row in rows[1:]:
        # Выравниваем строку по количеству колонок заголовка
        padded = row + [""] * (len(header) - len(row))
        lines.append("| " + " | ".join(str(c) for c in padded[:len(header)]) + " |")

    text = f"CSV-файл «{path.name}»:\n\n" + "\n".join(lines)
    return _truncate(text), None


def _read_spreadsheet(path: Path) -> Tuple[Optional[str], Optional[str]]:
    """Читает Excel / ODS и форматирует каждый лист как markdown-таблицу."""
    import openpyxl

    ext = path.suffix.lower()

    if ext in (".xlsx", ".xlsm"):
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        sheets_data = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = [[str(cell.value) if cell.value is not None else "" for cell in row] for row in ws.iter_rows()]
            sheets_data.append((sheet_name, rows))
        wb.close()

    elif ext == ".xls":
        import xlrd
        wb = xlrd.open_workbook(path)
        sheets_data = []
        for i in range(wb.nsheets):
            ws = wb.sheet_by_index(i)
            rows = [[str(ws.cell_value(r, c)) for c in range(ws.ncols)] for r in range(ws.nrows)]
            sheets_data.append((wb.sheet_names()[i], rows))

    elif ext == ".ods":
        from odf.opendocument import load as ods_load
        from odf.table import Table, TableRow, TableCell
        from odf.text import P

        doc = ods_load(path)
        sheets_data = []
        for sheet in doc.spreadsheet.getElementsByType(Table):
            rows = []
            for row_el in sheet.getElementsByType(TableRow):
                row = []
                for cell_el in row_el.getElementsByType(TableCell):
                    texts = cell_el.getElementsByType(P)
                    cell_text = " ".join(t.firstChild.data if t.firstChild else "" for t in texts)
                    row.append(cell_text)
                rows.append(row)
            sheets_data.append((sheet.getAttribute("name") or "Sheet", rows))

    else:
        return None, f"Неподдерживаемый формат таблицы: {ext}"

    # Форматируем каждый лист
    result_parts = []
    for sheet_name, rows in sheets_data:
        if not rows:
            continue
        header = rows[0]
        lines = [f"### Лист: {sheet_name}"]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        for row in rows[1:]:
            padded = row + [""] * (len(header) - len(row))
            lines.append("| " + " | ".join(padded[:len(header)]) + " |")
        result_parts.append("\n".join(lines))

    text = f"Таблица «{path.name}»:\n\n" + "\n\n".join(result_parts)
    return _truncate(text), None


def _read_word(path: Path) -> Tuple[Optional[str], Optional[str]]:
    """Читает DOCX / DOC / ODT и извлекает текст."""
    ext = path.suffix.lower()

    if ext == ".docx":
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Извлекаем таблицы
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                header = rows[0]
                lines = ["| " + " | ".join(header) + " |",
                         "| " + " | ".join(["---"] * len(header)) + " |"]
                for row in rows[1:]:
                    lines.append("| " + " | ".join(row[:len(header)]) + " |")
                table_texts.append("\n".join(lines))

        parts = paragraphs
        if table_texts:
            parts.append("\n\nТаблицы в документе:\n" + "\n\n".join(table_texts))
        text = f"Документ «{path.name}»:\n\n" + "\n".join(parts)
        return _truncate(text), None

    elif ext in (".doc",):
        # Попытка через antiword (если установлен) или textract
        try:
            import subprocess
            result = subprocess.run(["antiword", str(path)], capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                return _truncate(result.stdout), None
        except Exception:
            pass
        try:
            import textract
            text = textract.process(str(path)).decode("utf-8", errors="replace")
            return _truncate(f"Документ «{path.name}»:\n\n{text}"), None
        except Exception as e:
            return None, f"Не удалось прочитать .doc: {e}. Установите python-docx2txt или textract."

    elif ext == ".odt":
        try:
            from odf.opendocument import load as odt_load
            from odf.text import P
            doc = odt_load(path)
            texts = [str(p) for p in doc.getElementsByType(P)]
            text = f"Документ «{path.name}»:\n\n" + "\n".join(texts)
            return _truncate(text), None
        except Exception as e:
            return None, f"Не удалось прочитать .odt: {e}"

    return None, f"Неподдерживаемый формат документа: {ext}"


def _read_pdf(path: Path) -> Tuple[Optional[str], Optional[str]]:
    """Извлекает текст из PDF."""
    try:
        import pdfplumber
        texts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    texts.append(page_text)
                # Извлекаем таблицы со страницы
                for table in page.extract_tables():
                    if not table:
                        continue
                    header = [str(c) if c else "" for c in table[0]]
                    lines = ["| " + " | ".join(header) + " |",
                             "| " + " | ".join(["---"] * len(header)) + " |"]
                    for row in table[1:]:
                        cells = [str(c) if c else "" for c in row]
                        lines.append("| " + " | ".join(cells[:len(header)]) + " |")
                    texts.append("\n".join(lines))

        if not texts:
            return "(PDF не содержит извлекаемого текста — возможно, это сканированный документ)", None

        text = f"PDF «{path.name}»:\n\n" + "\n\n".join(texts)
        return _truncate(text), None

    except Exception as e:
        return None, f"Не удалось прочитать PDF: {e}"